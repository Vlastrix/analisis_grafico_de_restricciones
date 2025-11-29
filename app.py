"""
Aplicación Flask para Análisis Gráfico de Restricciones en Programación Lineal
Permite resolver y visualizar problemas de PL con 2 variables
"""

from flask import Flask, render_template, request, jsonify, send_from_directory, Response, send_file
import os
import re
import numpy as np
from scipy.optimize import linprog
from scipy.spatial import HalfspaceIntersection, ConvexHull
import json
from datetime import datetime
from io import BytesIO
import base64

# Librerías para exportación
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage
import plotly.graph_objects as go
import matplotlib
matplotlib.use('Agg')  # Backend sin interfaz gráfica
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon as MplPolygon

app = Flask(__name__)

def parse_constraint(constraint_str):
    """
    Parsea una restricción del formato: '2*x1 + 3*x2 <= 10'
    Retorna: coeficientes [a1, a2] y el valor b, y el tipo de restricción
    Lanza ValueError si la restricción es inválida
    """
    try:
        original_str = constraint_str
        
        # Validar que solo contenga x1 y x2 (no x3, x4, etc.)
        if re.search(r'x[3-9]|x\d\d+', constraint_str, re.IGNORECASE):
            raise ValueError(
                "El método gráfico solo funciona con 2 variables (x1 y x2). "
                "Encontró variables adicionales (x3, x4, etc.) en la restricción: " + original_str
            )
        
        # Reemplazar operadores
        constraint_str = constraint_str.replace('x1', 'x[0]').replace('x2', 'x[1]')
        constraint_str = constraint_str.replace('X1', 'x[0]').replace('X2', 'x[1]')
        
        # Identificar el tipo de restricción
        if '<=' in constraint_str:
            op = '<='
            left, right = constraint_str.split('<=')
        elif '>=' in constraint_str:
            op = '>='
            left, right = constraint_str.split('>=')
        elif '=' in constraint_str and '!' not in constraint_str:
            op = '='
            left, right = constraint_str.split('=')
        else:
            raise ValueError(
                "Operador no válido en restricción: " + original_str + 
                ". Use <=, >= o ="
            )
        
        # Evaluar el lado derecho (valor b)
        try:
            b = float(eval(right.strip()))
        except:
            raise ValueError(
                "El lado derecho de la restricción debe ser un número: " + original_str
            )
        
        # Extraer coeficientes del lado izquierdo
        left = left.strip()
        
        # Inicializar coeficientes
        a1, a2 = 0, 0
        
        # Método simple: evaluar en puntos específicos
        x = [1, 0]
        val1 = eval(left)
        x = [0, 1]
        val2 = eval(left)
        x = [0, 0]
        val0 = eval(left)
        
        a1 = val1 - val0
        a2 = val2 - val0
        
        # Validar que al menos un coeficiente sea no-cero
        if abs(a1) < 1e-10 and abs(a2) < 1e-10:
            raise ValueError(
                "La restricción no contiene variables x1 o x2: " + original_str
            )
        
        return [a1, a2], b, op
    
    except ValueError as e:
        # Re-lanzar errores de validación con mensaje claro
        raise e
    except Exception as e:
        raise ValueError(f"Error al parsear restricción '{original_str}': {str(e)}")

def find_feasible_region(constraints, bounds):
    """
    Encuentra los vértices de la región factible.
    Maneja regiones no acotadas añadiendo límites artificiales grandes para graficar.
    Usa método de intersecciones manuales para evitar problemas con Qhull.
    """
    if not constraints:
        return []
    
    # Convertir restricciones a formato de semiespacios
    A_ub = []
    b_ub = []
    has_x1_nonneg = False  # Bandera para x1 >= 0
    has_x2_nonneg = False  # Bandera para x2 >= 0
    
    for constraint in constraints:
        coeffs, b, op = constraint
        if op == '<=':
            A_ub.append(coeffs)
            b_ub.append(b)
        elif op == '>=':
            A_ub.append([-c for c in coeffs])
            b_ub.append(-b)
            # Detectar si ya existe x1 >= 0 o x2 >= 0
            if coeffs == [1, 0] and b == 0:
                has_x1_nonneg = True
            elif coeffs == [0, 1] and b == 0:
                has_x2_nonneg = True
    
    # Agregar restricciones de no negatividad SOLO si no fueron ingresadas por el usuario
    if not has_x1_nonneg:
        A_ub.append([-1, 0])  # -x1 <= 0 => x1 >= 0
        b_ub.append(0)
    if not has_x2_nonneg:
        A_ub.append([0, -1])  # -x2 <= 0 => x2 >= 0
        b_ub.append(0)
    
    # Agregar límites superiores si están definidos
    # Si no están definidos (región potencialmente no acotada), usar límites grandes para graficar
    x1_limit = bounds['x1_max'] if bounds['x1_max'] is not None else 1000
    x2_limit = bounds['x2_max'] if bounds['x2_max'] is not None else 1000
    
    A_ub.append([1, 0])
    b_ub.append(x1_limit)
    A_ub.append([0, 1])
    b_ub.append(x2_limit)
    
    A_ub = np.array(A_ub)
    b_ub = np.array(b_ub)
    
    # Usar método de intersecciones manuales (más robusto que HalfspaceIntersection)
    try:
        vertices = find_vertices_alternative(A_ub, b_ub)
        
        # Manejar casos degenerados
        if len(vertices) == 0:
            return []
        elif len(vertices) == 1:
            # Región es un solo punto
            return vertices
        elif len(vertices) == 2:
            # Región es un segmento (línea)
            return vertices
        else:
            # Región es un polígono: ordenar vértices usando ConvexHull
            try:
                vertices_array = np.array(vertices)
                hull = ConvexHull(vertices_array)
                vertices = vertices_array[hull.vertices].tolist()
            except Exception as e:
                # Si ConvexHull falla (puntos colineales), mantener vértices sin ordenar
                print(f"ConvexHull fallido: {e}")
                pass
            return vertices
    
    except Exception as e:
        print(f"Error al calcular región factible: {str(e)}")
        return []

def find_vertices_alternative(A_ub, b_ub):
    """
    Método alternativo para encontrar vértices de la región factible
    Encuentra intersecciones de pares de líneas
    """
    vertices = []
    n_constraints = len(A_ub)
    
    # Probar todas las combinaciones de 2 restricciones
    for i in range(n_constraints):
        for j in range(i+1, n_constraints):
            # Resolver sistema 2x2
            A = np.array([A_ub[i], A_ub[j]])
            b = np.array([b_ub[i], b_ub[j]])
            
            try:
                # Verificar que no sean paralelas
                det = np.linalg.det(A)
                if abs(det) < 1e-10:
                    continue
                
                # Resolver
                point = np.linalg.solve(A, b)
                
                # Verificar que el punto satisfaga todas las restricciones
                if np.all(A_ub @ point <= b_ub + 1e-6):
                    vertices.append(point.tolist())
            
            except:
                continue
    
    # Eliminar duplicados
    unique_vertices = []
    for v in vertices:
        is_duplicate = False
        for uv in unique_vertices:
            if np.allclose(v, uv, atol=1e-6):
                is_duplicate = True
                break
        if not is_duplicate:
            unique_vertices.append(v)
    
    return unique_vertices

def solve_lp(objective_coeffs, constraints, is_maximize, bounds):
    """
    Resuelve el problema de programación lineal.
    Retorna: (optimal_point, optimal_value, status)
    status puede ser: 'optimal', 'unbounded', 'infeasible', 'error'
    """
    if not constraints:
        return None, None, 'error'
    
    # Preparar coeficientes
    c = objective_coeffs if not is_maximize else [-x for x in objective_coeffs]
    
    # Preparar restricciones
    A_ub = []
    b_ub = []
    A_eq = []
    b_eq = []
    
    for constraint in constraints:
        coeffs, b, op = constraint
        if op == '<=':
            A_ub.append(coeffs)
            b_ub.append(b)
        elif op == '>=':
            A_ub.append([-x for x in coeffs])
            b_ub.append(-b)
        elif op == '=':
            A_eq.append(coeffs)
            b_eq.append(b)
    
    # Añadir límites superiores de x1_max y x2_max como restricciones
    # Solo si están especificados y no son None
    if bounds.get('x1_max') is not None:
        A_ub.append([1, 0])  # x1 <= x1_max
        b_ub.append(bounds['x1_max'])
    if bounds.get('x2_max') is not None:
        A_ub.append([0, 1])  # x2 <= x2_max
        b_ub.append(bounds['x2_max'])
    
    # Convertir a arrays numpy
    A_ub = np.array(A_ub) if A_ub else None
    b_ub = np.array(b_ub) if b_ub else None
    A_eq = np.array(A_eq) if A_eq else None
    b_eq = np.array(b_eq) if b_eq else None
    
    # Resolver
    try:
        result = linprog(
            c, 
            A_ub=A_ub, 
            b_ub=b_ub,
            A_eq=A_eq,
            b_eq=b_eq,
            bounds=(0, None),
            method='highs'
        )
        
        if result.success:
            optimal_value = -result.fun if is_maximize else result.fun
            return result.x.tolist(), optimal_value, 'optimal'
        elif result.status == 3:
            # Status 3 = unbounded
            return None, None, 'unbounded'
        elif result.status == 2:
            # Status 2 = infeasible
            return None, None, 'infeasible'
        else:
            return None, None, 'error'
    
    except Exception as e:
        print(f"Error al resolver LP: {str(e)}")
        return None, None, 'error'

def generate_solution_steps(data, result, format_type='default'):
    """
    Genera los pasos de resolución del problema de PL
    format_type: 'default' usa subscripts unicode, 'pdf' usa formato simple para compatibilidad
    """
    steps = []
    
    # Determinar formato de variables según el tipo
    if format_type == 'pdf':
        x1, x2 = 'x1', 'x2'
    else:
        x1, x2 = 'x₁', 'x₂'
    
    # Paso 1: Definición del problema
    obj_type = "Maximizar" if result.get('is_maximize') else "Minimizar"
    obj_coeffs = result.get('objective_coeffs', [0, 0])
    
    steps.append({
        'title': 'Paso 1: Definición del Problema',
        'content': f"{obj_type} Z = {obj_coeffs[0]}{x1} + {obj_coeffs[1]}{x2}"
    })
    
    # Paso 2: Restricciones
    constraints_text = "Sujeto a:\n"
    for i, constraint in enumerate(result.get('constraint_lines', []), 1):
        coeffs = constraint['coeffs']
        b = constraint['b']
        op = constraint['op']
        constraints_text += f"  {i}. {coeffs[0]}{x1} + {coeffs[1]}{x2} {op} {b}\n"
    constraints_text += f"  {x1}, {x2} >= 0 (No negatividad)"
    
    steps.append({
        'title': 'Paso 2: Restricciones del Problema',
        'content': constraints_text
    })
    
    # Paso 3: Región factible
    vertices = result.get('vertices', [])
    vertices_text = f"Se identificaron {len(vertices)} vértices de la región factible:\n"
    for i, v in enumerate(vertices, 1):
        vertices_text += f"  V{i} = ({v[0]:.4f}, {v[1]:.4f})\n"
    
    steps.append({
        'title': 'Paso 3: Identificación de Vértices',
        'content': vertices_text
    })
    
    # Paso 4: Evaluación de la función objetivo
    if result.get('success') and result.get('optimal_value') is not None:
        eval_text = "Evaluación de Z en cada vértice:\n"
        for i, v in enumerate(vertices, 1):
            z_value = obj_coeffs[0] * v[0] + obj_coeffs[1] * v[1]
            eval_text += f"  Z(V{i}) = {obj_coeffs[0]}({v[0]:.4f}) + {obj_coeffs[1]}({v[1]:.4f}) = {z_value:.4f}\n"
        
        steps.append({
            'title': 'Paso 4: Evaluación de la Función Objetivo',
            'content': eval_text
        })
        
        # Paso 5: Solución óptima
        optimal_point = result.get('optimal_point', [0, 0])
        optimal_value = result.get('optimal_value', 0)
        solution_text = f"Punto óptimo: ({optimal_point[0]:.4f}, {optimal_point[1]:.4f})\n"
        solution_text += f"Valor óptimo: Z* = {optimal_value:.4f}"
        
        steps.append({
            'title': 'Paso 5: Solución Óptima',
            'content': solution_text
        })
    
    return steps

def create_plotly_image(result, x1_max=10, x2_max=10):
    """
    Crea una imagen del gráfico usando Plotly
    """
    try:
        fig = go.Figure()
        
        # Región factible
        if result.get('vertices'):
            vertices = result['vertices']
            if len(vertices) >= 3:
                # Cerrar el polígono
                vertices_closed = vertices + [vertices[0]]
                x_coords = [v[0] for v in vertices_closed]
                y_coords = [v[1] for v in vertices_closed]
                
                fig.add_trace(go.Scatter(
                    x=x_coords,
                    y=y_coords,
                    fill='toself',
                    fillcolor='rgba(135, 206, 250, 0.3)',
                    line=dict(color='blue', width=2),
                    name='Región Factible'
                ))
            
            # Vértices
            x_vert = [v[0] for v in vertices]
            y_vert = [v[1] for v in vertices]
            fig.add_trace(go.Scatter(
                x=x_vert,
                y=y_vert,
                mode='markers',
                marker=dict(size=10, color='blue'),
                name='Vértices'
            ))
        
        # Punto óptimo
        if result.get('optimal_point'):
            opt = result['optimal_point']
            fig.add_trace(go.Scatter(
                x=[opt[0]],
                y=[opt[1]],
                mode='markers',
                marker=dict(size=15, color='red', symbol='star'),
                name=f'Óptimo: ({opt[0]:.2f}, {opt[1]:.2f})'
            ))
        
        # Líneas de restricciones
        for constraint in result.get('constraint_lines', []):
            coeffs = constraint['coeffs']
            b = constraint['b']
            
            if coeffs[1] != 0:
                x_line = np.linspace(0, x1_max, 100)
                y_line = (b - coeffs[0] * x_line) / coeffs[1]
                
                fig.add_trace(go.Scatter(
                    x=x_line,
                    y=y_line,
                    mode='lines',
                    line=dict(dash='dash', width=1),
                    name=constraint.get('label', f"Restricción {coeffs}")
                ))
        
        fig.update_layout(
            title='Análisis Gráfico del Problema de PL',
            xaxis_title='x₁',
            yaxis_title='x₂',
            xaxis=dict(range=[0, x1_max]),
            yaxis=dict(range=[0, x2_max]),
            showlegend=True,
            width=800,
            height=600
        )
        
        # Convertir a imagen - intentar con kaleido
        print("Generando imagen con Plotly...")
        
        try:
            img_bytes = fig.to_image(format="png", engine="kaleido")
            print("Imagen generada exitosamente con kaleido")
            return img_bytes
        except Exception as e:
            print(f"Kaleido falló: {e}")
            # Si kaleido falla, retornar None y continuar sin imagen
            return None
            
    except Exception as e:
        print(f"Error al generar imagen de Plotly: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_matplotlib_image(result, x1_max=10, x2_max=10):
    """
    Crea una imagen del gráfico usando Matplotlib (más confiable que Kaleido)
    """
    try:
        print("Generando imagen con Matplotlib...")
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Colores
        colors_list = ['#ef4444', '#f59e0b', '#10b981', '#3b82f6', '#8b5cf6', '#ec4899', '#14b8a6', '#f97316']
        
        # Dibujar líneas de restricciones
        for idx, constraint in enumerate(result.get('constraint_lines', [])):
            coeffs = constraint['coeffs']
            b = constraint['b']
            color = colors_list[idx % len(colors_list)]
            
            if coeffs[1] != 0:
                x_line = np.linspace(0, x1_max, 100)
                y_line = (b - coeffs[0] * x_line) / coeffs[1]
                # Filtrar valores dentro del rango
                mask = (y_line >= 0) & (y_line <= x2_max)
                ax.plot(x_line[mask], y_line[mask], '--', color=color, linewidth=2, 
                       label=constraint.get('label', f"Restricción {idx+1}"), alpha=0.7)
            elif coeffs[0] != 0:
                # Línea vertical
                x_val = b / coeffs[0]
                if 0 <= x_val <= x1_max:
                    ax.axvline(x=x_val, color=color, linestyle='--', linewidth=2,
                             label=constraint.get('label', f"Restricción {idx+1}"), alpha=0.7)
        
        # Dibujar región factible
        if result.get('vertices') and len(result['vertices']) >= 3:
            vertices = result['vertices']
            # Ordenar vértices
            vertices_array = np.array(vertices)
            # Calcular centroide
            cx = np.mean(vertices_array[:, 0])
            cy = np.mean(vertices_array[:, 1])
            # Ordenar por ángulo
            angles = np.arctan2(vertices_array[:, 1] - cy, vertices_array[:, 0] - cx)
            sorted_indices = np.argsort(angles)
            sorted_vertices = vertices_array[sorted_indices]
            
            # Crear polígono
            polygon = MplPolygon(sorted_vertices, alpha=0.3, facecolor='lightblue', 
                                edgecolor='blue', linewidth=2, label='Región Factible')
            ax.add_patch(polygon)
            
            # Dibujar vértices
            ax.scatter(sorted_vertices[:, 0], sorted_vertices[:, 1], 
                      color='blue', s=100, zorder=5, label='Vértices', marker='o')
        
        # Dibujar punto óptimo
        if result.get('optimal_point'):
            opt = result['optimal_point']
            ax.scatter([opt[0]], [opt[1]], color='red', s=300, zorder=10,
                      label=f'Óptimo ({opt[0]:.2f}, {opt[1]:.2f})', marker='*')
        
        # Configurar ejes
        ax.set_xlim(0, x1_max)
        ax.set_ylim(0, x2_max)
        ax.set_xlabel('x₁', fontsize=12, fontweight='bold')
        ax.set_ylabel('x₂', fontsize=12, fontweight='bold')
        ax.set_title('Análisis Gráfico del Problema de PL', fontsize=14, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.legend(loc='best', framealpha=0.9)
        
        # Guardar en memoria
        img_stream = BytesIO()
        plt.tight_layout()
        plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        img_stream.seek(0)
        
        print("Imagen generada exitosamente con Matplotlib")
        return img_stream.read()
        
    except Exception as e:
        print(f"Error al generar imagen con Matplotlib: {e}")
        import traceback
        traceback.print_exc()
        return None

def export_to_word(data, result):
    """
    Exporta el problema y su solución a un documento Word
    """
    try:
        doc = Document()
        
        # Título
        title = doc.add_heading('Análisis de Programación Lineal', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        date_para = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Generar pasos
        steps = generate_solution_steps(data, result, format_type='default')
        
        # Agregar cada paso
        for step in steps:
            doc.add_heading(step['title'], level=1)
            
            # Agregar contenido con formato
            for line in step['content'].split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    if line.strip().startswith(('V', 'Z(')):
                        p.style = 'List Bullet'
        
        # Agregar gráfico si está disponible
        try:
            x1_max = data.get('x1_max', 10)
            x2_max = data.get('x2_max', 10)
            
            # Usar SOLO Matplotlib (más confiable que Kaleido en Windows)
            print("Generando imagen con Matplotlib...")
            img_bytes = create_matplotlib_image(result, x1_max, x2_max)
            
            if img_bytes:
                doc.add_paragraph()  # Espacio en lugar de página nueva
                doc.add_heading('Visualización Gráfica', level=1)
                
                # Guardar temporalmente la imagen
                img_stream = BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
                print("Imagen agregada exitosamente a Word")
            else:
                print("No se pudo generar la imagen para Word")
                doc.add_paragraph()  # Espacio en lugar de página nueva
                doc.add_heading('Visualización Gráfica', level=1)
                doc.add_paragraph('Nota: El gráfico no pudo ser generado automáticamente. ')
                doc.add_paragraph('Por favor, visualiza el gráfico en la aplicación web.')
        except Exception as e:
            print(f"Error al agregar gráfico a Word: {e}")
            import traceback
            traceback.print_exc()
            doc.add_paragraph()  # Espacio en lugar de página nueva
            doc.add_heading('Visualización Gráfica', level=1)
            doc.add_paragraph('Nota: El gráfico no pudo ser generado automáticamente. ')
            doc.add_paragraph('Por favor, visualiza el gráfico en la aplicación web.')
        
        # Guardar en memoria
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    except Exception as e:
        print(f"Error general en export_to_word: {e}")
        import traceback
        traceback.print_exc()
        raise

def export_to_pdf(data, result):
    """
    Exporta el problema y su solución a PDF
    """
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Estilo personalizado
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=1  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#3498db'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Título
        elements.append(Paragraph('Análisis de Programación Lineal', title_style))
        elements.append(Paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}', styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # Generar pasos con formato PDF (sin unicode especial)
        steps = generate_solution_steps(data, result, format_type='pdf')
        
        # Agregar cada paso
        for step in steps:
            elements.append(Paragraph(step['title'], heading_style))
            
            # Agregar contenido
            for line in step['content'].split('\n'):
                if line.strip():
                    # Escapar caracteres especiales para XML
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    elements.append(Paragraph(safe_line, styles['Normal']))
                    elements.append(Spacer(1, 6))
            
            elements.append(Spacer(1, 12))
        
        # Agregar gráfico
        try:
            x1_max = data.get('x1_max', 10)
            x2_max = data.get('x2_max', 10)
            
            # Usar SOLO Matplotlib (más confiable que Kaleido en Windows)
            print("Generando imagen con Matplotlib...")
            img_bytes = create_matplotlib_image(result, x1_max, x2_max)
            
            if img_bytes:
                img_stream = BytesIO(img_bytes)
                img = RLImage(img_stream, width=5*inch, height=3.75*inch)
                elements.append(Spacer(1, 20))
                elements.append(Paragraph('Visualización Gráfica', heading_style))
                elements.append(img)
                print("Imagen agregada exitosamente a PDF")
            else:
                print("No se pudo generar la imagen para PDF")
                elements.append(Spacer(1, 20))
                elements.append(Paragraph('Visualización Gráfica', heading_style))
                elements.append(Paragraph('Nota: El gráfico no pudo ser generado automáticamente. Por favor, visualiza el gráfico en la aplicación web.', styles['Normal']))
        except Exception as e:
            print(f"Error al agregar gráfico al PDF: {e}")
            import traceback
            traceback.print_exc()
            elements.append(Spacer(1, 20))
            elements.append(Paragraph('Visualización Gráfica', heading_style))
            elements.append(Paragraph('Nota: El gráfico no pudo ser generado automáticamente. Por favor, visualiza el gráfico en la aplicación web.', styles['Normal']))
        
        doc.build(elements)
        buffer.seek(0)
        return buffer
    except Exception as e:
        print(f"Error general en export_to_pdf: {e}")
        import traceback
        traceback.print_exc()
        raise

def export_to_excel(data, result):
    """
    Exporta el problema y su solución a Excel
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Análisis PL"
    
    # Estilos
    header_fill = PatternFill(start_color="3498db", end_color="3498db", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=12)
    title_font = Font(bold=True, size=16, color="2c3e50")
    
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    row = 1
    
    # Título
    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = 'ANÁLISIS DE PROGRAMACIÓN LINEAL'
    cell.font = title_font
    cell.alignment = Alignment(horizontal='center', vertical='center')
    row += 1
    
    # Fecha
    ws.merge_cells(f'A{row}:D{row}')
    ws[f'A{row}'] = f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws[f'A{row}'].alignment = Alignment(horizontal='right')
    row += 2
    
    # Generar pasos con formato simple
    steps = generate_solution_steps(data, result, format_type='default')
    
    for step in steps:
        # Título del paso
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws[f'A{row}']
        cell.value = step['title']
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='left', vertical='center')
        cell.border = border
        row += 1
        
        # Contenido del paso
        for line in step['content'].split('\n'):
            if line.strip():
                ws.merge_cells(f'A{row}:D{row}')
                ws[f'A{row}'] = line
                ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                row += 1
        
        row += 1
    
    # Tabla de vértices
    if result.get('vertices'):
        row += 1
        ws[f'A{row}'] = 'Vértice'
        ws[f'B{row}'] = 'x₁'
        ws[f'C{row}'] = 'x₂'
        ws[f'D{row}'] = 'Z'
        
        for col in ['A', 'B', 'C', 'D']:
            cell = ws[f'{col}{row}']
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal='center')
        
        row += 1
        
        obj_coeffs = result.get('objective_coeffs', [0, 0])
        for i, v in enumerate(result['vertices'], 1):
            z_value = obj_coeffs[0] * v[0] + obj_coeffs[1] * v[1]
            ws[f'A{row}'] = f'V{i}'
            ws[f'B{row}'] = round(v[0], 4)
            ws[f'C{row}'] = round(v[1], 4)
            ws[f'D{row}'] = round(z_value, 4)
            
            for col in ['A', 'B', 'C', 'D']:
                ws[f'{col}{row}'].border = border
                ws[f'{col}{row}'].alignment = Alignment(horizontal='center')
            
            row += 1
    
    # Ajustar anchos de columna
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    
    # Agregar gráfico como imagen
    try:
        row += 2
        x1_max = data.get('x1_max', 10)
        x2_max = data.get('x2_max', 10)
        
        # Título del gráfico
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws[f'A{row}']
        cell.value = 'Visualización Gráfica'
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='left', vertical='center')
        cell.border = border
        row += 1
        
        # Generar imagen con Matplotlib
        print("Generando imagen con Matplotlib para Excel...")
        img_bytes = create_matplotlib_image(result, x1_max, x2_max)
        
        if img_bytes:
            # Guardar imagen temporalmente en memoria
            img_stream = BytesIO(img_bytes)
            
            # Crear objeto de imagen para Excel
            img = XLImage(img_stream)
            
            # Ajustar tamaño de la imagen (ancho en píxeles)
            img.width = 600
            img.height = 450
            
            # Insertar imagen en la celda A{row}
            ws.add_image(img, f'A{row}')
            print("Imagen agregada exitosamente a Excel")
            
            # Ajustar altura de las filas para que la imagen se vea bien
            # Aproximadamente 450 píxeles / 1.33 = ~338 puntos, dividido en múltiples filas
            for i in range(30):  # Espacio para ~30 filas de altura
                ws.row_dimensions[row + i].height = 15
        else:
            print("No se pudo generar la imagen para Excel")
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = 'Nota: El gráfico no pudo ser generado. Por favor, visualiza el gráfico en la aplicación web.'
            ws[f'A{row}'].alignment = Alignment(horizontal='left', wrap_text=True)
    except Exception as e:
        print(f"Error al agregar gráfico a Excel: {e}")
        import traceback
        traceback.print_exc()
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'] = 'Nota: El gráfico no pudo ser generado. Por favor, visualiza el gráfico en la aplicación web.'
        ws[f'A{row}'].alignment = Alignment(horizontal='left', wrap_text=True)
    
    # Guardar en memoria
    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

@app.route('/')
def index():
    """Página principal"""
    return render_template('index.html')

@app.route('/solve', methods=['POST'])
def solve():
    """Endpoint para resolver el problema de PL con validaciones completas"""
    try:
        data = request.get_json(force=True)

        # ===== VALIDACIONES BACKEND =====
        
        # --- Parsear función objetivo (soporta dos formatos) ---
        # Formato estructurado: { "objective": { "sense": "max", "coefficients": [3,4] } }
        # Formato antiguo: campos individuales obj_x1 / obj_x2 y objective_type
        obj_coeffs = [0.0, 0.0]
        is_maximize = True
        if isinstance(data.get('objective'), dict):
            obj = data.get('objective', {})
            coeffs = obj.get('coefficients') or obj.get('coeff') or []
            if len(coeffs) >= 2:
                obj_coeffs = [float(coeffs[0]), float(coeffs[1])]
            sense = obj.get('sense') or obj.get('type')
            if sense:
                is_maximize = str(sense).lower() in ('max', 'maximize')
        else:
            # Fallback al formato por inputs individuales
            obj_coeffs = [
                float(data.get('obj_x1', 0)),
                float(data.get('obj_x2', 0))
            ]
            is_maximize = data.get('objective_type', 'maximize') == 'maximize'

        # 1. Validar función objetivo
        if obj_coeffs[0] == 0 and obj_coeffs[1] == 0:
            return jsonify({
                'success': False,
                'error': 'La función objetivo no puede tener todos los coeficientes en cero. Debe especificar al menos un coeficiente no nulo para x1 o x2.'
            })
        
        if not isinstance(obj_coeffs[0], (int, float)) or not isinstance(obj_coeffs[1], (int, float)):
            return jsonify({
                'success': False,
                'error': 'Los coeficientes de la función objetivo deben ser números válidos.'
            })

        # --- Parsear restricciones (soporta cadenas y objetos) ---
        constraints = []
        constraint_inputs = data.get('constraints', [])
        
        # 2. Validar que haya restricciones
        if not constraint_inputs or len(constraint_inputs) == 0:
            return jsonify({
                'success': False,
                'error': 'Debe ingresar al menos una restricción para resolver el problema de programación lineal.'
            })
        
        # 3. Validar número máximo de restricciones (límite razonable)
        if len(constraint_inputs) > 20:
            return jsonify({
                'success': False,
                'error': f'Demasiadas restricciones ({len(constraint_inputs)}). El método gráfico funciona mejor con 20 o menos restricciones.'
            })

        for idx, item in enumerate(constraint_inputs, 1):
            # Si llega un objeto con coeficientes y rhs
            if isinstance(item, dict):
                try:
                    coeffs = item.get('coefficients') or item.get('coeff')
                    b = item.get('rhs') if 'rhs' in item else item.get('b')
                    op = item.get('op') if 'op' in item else item.get('operator', '<=')
                    
                    if coeffs is None or b is None:
                        return jsonify({
                            'success': False,
                            'error': f"Restricción {idx} incompleta: faltan coeficientes o valor del lado derecho."
                        })
                    
                    # Validar que solo haya 2 coeficientes
                    if len(coeffs) != 2:
                        return jsonify({
                            'success': False,
                            'error': f"Restricción {idx}: El método gráfico solo funciona con 2 variables (x1 y x2). Encontró {len(coeffs)} variables."
                        })
                    
                    constraints.append(([float(coeffs[0]), float(coeffs[1])], float(b), str(op)))
                except ValueError as ve:
                    return jsonify({'success': False, 'error': f"Restricción {idx}: {str(ve)}"})
                except Exception as e:
                    return jsonify({'success': False, 'error': f"Error en restricción {idx}: {str(e)}"})
            else:
                # Texto: "2*x1 + x2 <= 10"
                try:
                    if isinstance(item, str) and item.strip():
                        parsed = parse_constraint(item)
                        constraints.append(parsed)
                    elif not item or (isinstance(item, str) and not item.strip()):
                        return jsonify({
                            'success': False,
                            'error': f"Restricción {idx} está vacía. Por favor, elimínela o complétela."
                        })
                except ValueError as ve:
                    # Error específico de validación
                    return jsonify({
                        'success': False,
                        'error': f"Restricción {idx}: {str(ve)}"
                    })
                except Exception as e:
                    return jsonify({
                        'success': False,
                        'error': f"Error al procesar restricción {idx} ('{item}'): {str(e)}"
                    })

        # Obtener límites para el gráfico (soporta campo 'bounds' o x1_max/x2_max)
        x1_max = None
        x2_max = None
        if isinstance(data.get('bounds'), list) and len(data.get('bounds')) >= 2:
            # bounds: [[0, null], [0, null]] or [x1_max, x2_max]
            b = data.get('bounds')
            # If it's a pair of ranges, try to extract upper bounds if present
            try:
                # If it's like [[0,null],[0,null]] then no upper bounds
                if all(isinstance(x, list) for x in b):
                    # try to read second element of each bound if numeric
                    x1_upper = b[0][1]
                    x2_upper = b[1][1]
                    x1_max = None if x1_upper is None else float(x1_upper)
                    x2_max = None if x2_upper is None else float(x2_upper)
                else:
                    # simple list [x1_max, x2_max]
                    x1_max = None if b[0] is None else float(b[0])
                    x2_max = None if b[1] is None else float(b[1])
            except Exception:
                x1_max = None
                x2_max = None

        # fallback to explicit fields or defaults
        if x1_max is None:
            x1_max = data.get('x1_max', 20)
        if x2_max is None:
            x2_max = data.get('x2_max', 20)
        
        # 4. Validar límites del gráfico
        try:
            x1_max = float(x1_max)
            x2_max = float(x2_max)
            
            if x1_max <= 0 or x2_max <= 0:
                return jsonify({
                    'success': False,
                    'error': 'Los límites del gráfico (x₁ max y x₂ max) deben ser mayores que cero.'
                })
            
            if x1_max > 100000 or x2_max > 100000:
                return jsonify({
                    'success': False,
                    'error': 'Los límites del gráfico son demasiado grandes. Use valores menores a 100,000 para mejor visualización.'
                })
        except (ValueError, TypeError):
            return jsonify({
                'success': False,
                'error': 'Los límites del gráfico deben ser números válidos.'
            })

        bounds = {
            'x1_max': None if x1_max is None else float(x1_max),
            'x2_max': None if x2_max is None else float(x2_max)
        }
        
        # Encontrar región factible
        vertices = find_feasible_region(constraints, bounds)
        
        # Filtrar restricciones redundantes de no-negatividad (x1>=0, x2>=0)
        # porque linprog ya las asume con bounds=(0,None)
        constraints_for_solver = []
        for coeffs, b, op in constraints:
            # Filtrar x1 >= 0: coeffs=[1,0], b=0, op='>='
            if coeffs == [1, 0] and b == 0 and op == '>=':
                continue  # Saltar, ya está en bounds
            # Filtrar x2 >= 0: coeffs=[0,1], b=0, op='>='
            if coeffs == [0, 1] and b == 0 and op == '>=':
                continue  # Saltar, ya está en bounds
            constraints_for_solver.append((coeffs, b, op))
        
        # Resolver problema de optimización
        optimal_point, optimal_value, status = solve_lp(obj_coeffs, constraints_for_solver, is_maximize, bounds)

        # Manejar caso de problema no acotado
        if status == 'unbounded':
            return jsonify({
                'success': False,
                'unbounded': True,
                'error': 'Problema no acotado: la región factible es infinita y la función objetivo puede crecer sin límite.' if is_maximize else 'Problema no acotado: la región factible es infinita y la función objetivo puede decrecer sin límite.',
                'vertices': vertices,
                'constraint_lines': [
                    {
                        'coeffs': coeffs,
                        'b': b,
                        'op': op,
                        'label': (orig.strip() if isinstance(orig, str) else orig.get('label') or f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}") if isinstance(orig, (str, dict)) else f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}"
                    }
                    for orig, (coeffs, b, op) in zip(constraint_inputs, constraints)
                ],
                'objective_coeffs': obj_coeffs,
                'is_maximize': is_maximize
            })
        
        # Manejar caso infactible
        if status == 'infeasible':
            return jsonify({
                'success': False,
                'infeasible': True,
                'error': 'Problema infactible: no existe ningún punto que satisfaga todas las restricciones.',
                'constraint_lines': [
                    {
                        'coeffs': coeffs,
                        'b': b,
                        'op': op,
                        'label': (orig.strip() if isinstance(orig, str) else orig.get('label') or f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}") if isinstance(orig, (str, dict)) else f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}"
                    }
                    for orig, (coeffs, b, op) in zip(constraint_inputs, constraints)
                ],
                'objective_coeffs': obj_coeffs,
                'is_maximize': is_maximize
            })

        # Identificar todos los puntos óptimos entre los vértices (si hay múltiples soluciones)
        optimal_points = []
        try:
            if vertices and optimal_value is not None:
                tol = 1e-6
                for v in vertices:
                    val = float(obj_coeffs[0]) * float(v[0]) + float(obj_coeffs[1]) * float(v[1])
                    if abs(val - optimal_value) <= tol:
                        optimal_points.append([float(v[0]), float(v[1])])
        except Exception:
            optimal_points = []
        
        # Preparar datos para graficar
        constraint_lines = []
        # Build labels depending on original input (string or structured)
        for orig, (coeffs, b, op) in zip(constraint_inputs, constraints):
            if isinstance(orig, str):
                label = orig.strip()
            elif isinstance(orig, dict):
                # try to build a human label
                label = orig.get('label') or f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}"
            else:
                label = f"{coeffs[0]}*x1 + {coeffs[1]}*x2 {op} {b}"

            constraint_lines.append({
                'coeffs': coeffs,
                'b': b,
                'op': op,
                'label': label
            })
        
        return jsonify({
            'success': True,
            'vertices': vertices,
            'optimal_point': optimal_point,
            'optimal_points': optimal_points,
            'optimal_value': optimal_value,
            'constraint_lines': constraint_lines,
            'objective_coeffs': obj_coeffs,
            'is_maximize': is_maximize
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/export/word', methods=['POST'])
def export_word():
    """Endpoint para exportar a Word"""
    try:
        print("=== Iniciando exportación a Word ===")
        data = request.get_json()
        print(f"Datos recibidos: {data.keys() if data else 'None'}")
        
        result = data.get('result', {})
        original_data = data.get('data', {})
        
        print("Llamando a export_to_word...")
        file_stream = export_to_word(original_data, result)
        print("Archivo Word generado exitosamente")
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'analisis_pl_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
    except Exception as e:
        print(f"ERROR en export_word: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    """Endpoint para exportar a PDF"""
    try:
        print("=== Iniciando exportación a PDF ===")
        data = request.get_json()
        print(f"Datos recibidos: {data.keys() if data else 'None'}")
        
        result = data.get('result', {})
        original_data = data.get('data', {})
        
        print("Llamando a export_to_pdf...")
        file_stream = export_to_pdf(original_data, result)
        print("Archivo PDF generado exitosamente")
        
        return send_file(
            file_stream,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'analisis_pl_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        )
    except Exception as e:
        print(f"ERROR en export_pdf: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500
        
        return send_file(
            file_stream,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'analisis_pl_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/export/excel', methods=['POST'])
def export_excel():
    """Endpoint para exportar a Excel"""
    try:
        data = request.get_json()
        result = data.get('result', {})
        original_data = data.get('data', {})
        
        file_stream = export_to_excel(original_data, result)
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'analisis_pl_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)


@app.route('/favicon.ico')
def favicon():
    """Serve favicon if present in static, otherwise return 204 so browser doesn't log a 404."""
    ico_path = os.path.join(app.root_path, 'static', 'favicon.ico')
    if os.path.exists(ico_path):
        return send_from_directory(os.path.join(app.root_path, 'static'), 'favicon.ico')
    return Response(status=204)
